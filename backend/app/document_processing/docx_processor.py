# backend/app/document_processing/docx_processor.py
from app.document_processing.base import DocumentResult


class DocxProcessor:
    def process(self, file_path: str) -> DocumentResult:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        return DocumentResult(
            text=text,
            page_count=1,
            extraction_method="python-docx",
        )
